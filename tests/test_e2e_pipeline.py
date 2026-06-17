import argparse
import json
from pathlib import Path

from docx import Document

from src.load import collect_errors, write_json, write_processing_log
from src.main import run_pipeline


def _create_docx(path: Path) -> None:
    document = Document()
    document.core_properties.title = "DOCX E2E"
    document.core_properties.author = "Test Author"
    document.add_heading("DOCX E2E", level=1)
    document.add_paragraph("Автор: Test Author")
    document.add_paragraph("SQL таблица индекс транзакция query database.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Термин"
    table.cell(0, 1).text = "Описание"
    table.cell(1, 0).text = "index"
    table.cell(1, 1).text = "ускоряет поиск"
    document.save(path)


def test_pipeline_writes_outputs_and_new_extract_fields(tmp_path: Path) -> None:
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    output_dir.mkdir()

    html = input_dir / "sample.html"
    html.write_text(
        """
        <html>
          <head><title>HTML E2E</title><meta name="author" content="HTML Author"></head>
          <body><h1>HTML E2E</h1><h2>Алгоритм</h2><p>Алгоритм граф дерево поиск.</p></body>
        </html>
        """,
        encoding="utf-8",
    )
    _create_docx(input_dir / "sample.docx")

    args = argparse.Namespace(
        input=str(input_dir),
        output=str(output_dir / "dataset.json"),
        metrics_output=str(output_dir / "metrics.json"),
        log_output=str(output_dir / "processing_log.csv"),
        errors_output=str(output_dir / "errors.json"),
        pretty=True,
        limit=None,
    )

    dataset, metrics, log_rows = run_pipeline(args)
    write_json(dataset, args.output, pretty=True)
    write_json(metrics, args.metrics_output, pretty=True)
    write_json(collect_errors(dataset["documents"]), args.errors_output, pretty=True)
    write_processing_log(log_rows, args.log_output)

    for file_name in ["dataset.json", "metrics.json", "processing_log.csv", "errors.json"]:
        assert (output_dir / file_name).exists()

    saved_dataset = json.loads((output_dir / "dataset.json").read_text(encoding="utf-8"))
    assert saved_dataset["processed_files"] == 2
    assert saved_dataset["failed_files"] == 0
    assert metrics["metadata_completeness_common"] > 0
    assert "metadata_completeness_format_specific" in metrics
    assert "metadata_completeness_total" in metrics

    for document in saved_dataset["documents"]:
        assert "extractor_used" in document
        assert "fallback_used" in document
        assert "headings" in document
        assert "embedded_metadata" in document
        assert "ingest_status" in document
        assert "extractor_status" in document
