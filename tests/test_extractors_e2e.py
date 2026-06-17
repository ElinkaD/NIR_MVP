from pathlib import Path

import pytest
from docx import Document

from src.extract.docx_extractor import DocxExtractor
from src.extract.html_extractor import HtmlExtractor
from src.extract.pdf_extractor import PdfExtractor


def test_html_extractor_e2e(tmp_path: Path) -> None:
    path = tmp_path / "page.html"
    path.write_text(
        """
        <html>
          <head><title>Алгоритмы</title><meta name="description" content="graph algorithms"></head>
          <body><nav>remove me</nav><h1>Граф</h1><h2>Поиск</h2><p>Алгоритм поиска в графе.</p></body>
        </html>
        """,
        encoding="utf-8",
    )

    result = HtmlExtractor().extract(path)

    assert result.extraction_status == "success"
    assert result.extractor_used == "BeautifulSoup"
    assert result.fallback_used is False
    assert result.title == "Алгоритмы"
    assert result.headings == ["Граф", "Поиск"]
    assert result.metadata["meta"]["description"] == "graph algorithms"
    assert "remove me" not in result.text


def test_docx_extractor_e2e(tmp_path: Path) -> None:
    path = tmp_path / "lecture.docx"
    document = Document()
    document.core_properties.title = "Репликация"
    document.core_properties.author = "Test Author"
    document.add_heading("Репликация", level=1)
    document.add_paragraph("Репликация копирует данные между узлами.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "leader"
    table.cell(0, 1).text = "follower"
    document.save(path)

    result = DocxExtractor().extract(path)

    assert result.extraction_status == "success"
    assert result.extractor_used == "python-docx"
    assert result.title == "Репликация"
    assert result.author == "Test Author"
    assert result.headings == ["Репликация"]
    assert "leader | follower" in result.text


def test_pdf_extractor_e2e_optional_text_pdf(tmp_path: Path) -> None:
    reportlab = pytest.importorskip("reportlab.pdfgen.canvas")
    path = tmp_path / "sample.pdf"
    canvas = reportlab.Canvas(str(path))
    canvas.setTitle("PDF E2E")
    canvas.drawString(100, 750, "Direct Memory Access DMA")
    canvas.save()

    result = PdfExtractor().extract(path)

    assert result.extraction_status == "success"
    assert result.extractor_used in {"pdfplumber", "PyPDF2"}
    assert "DMA" in result.text
    assert result.page_count == 1
