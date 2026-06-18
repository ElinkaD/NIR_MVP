"""Create synthetic DOCX educational samples for extractor testing."""

from __future__ import annotations

import argparse
import csv
from pathlib import Path

from docx import Document

MANIFEST_FIELDS = [
    "file_id",
    "file_name",
    "relative_path",
    "format",
    "source_type",
    "source_url",
    "topic_expected",
    "difficulty_expected",
    "has_text_layer",
    "is_duplicate",
    "is_bad_file",
    "comment",
]

SAMPLES = [
    {
        "slug": "python",
        "title": "Основы Python",
        "author": "Synthetic NIR Dataset",
        "topic": "programming",
        "difficulty": "easy",
        "paragraphs": [
            "Автор: Synthetic NIR Dataset",
            "Python используется для описания алгоритмов, функций, классов и обработки данных.",
            "Пример кода: def train(model, dataset): return model.fit(dataset).",
        ],
        "table": [["Термин", "Описание"], ["function", "переиспользуемый блок кода"], ["class", "шаблон объекта"]],
    },
    {
        "slug": "sql",
        "title": "SQL и индексы",
        "author": "Synthetic NIR Dataset",
        "topic": "databases",
        "difficulty": "medium",
        "paragraphs": [
            "SQL-запрос использует таблицу, индекс и транзакцию.",
            "SELECT помогает получить данные, а индекс ускоряет поиск строк.",
        ],
        "table": [["Операция", "Назначение"], ["SELECT", "чтение"], ["INSERT", "добавление"]],
    },
    {
        "slug": "ml",
        "title": "Машинное обучение",
        "author": "",
        "topic": "machine_learning",
        "difficulty": "medium",
        "paragraphs": [
            "Модель обучается на тренировочном наборе, а качество проверяется на тестовой выборке.",
            "Градиентный спуск минимизирует ошибку y = f(x).",
        ],
        "table": [["Метод", "Задача"], ["classification", "классификация"], ["regression", "регрессия"]],
    },
    {
        "slug": "os",
        "title": "Операционные системы",
        "author": "Synthetic NIR Dataset",
        "topic": "operating_systems",
        "difficulty": "medium",
        "paragraphs": [
            "Процесс, поток, память и ядро являются базовыми понятиями операционной системы.",
            "Планировщик распределяет процессорное время между потоками.",
        ],
        "table": [["Понятие", "Пример"], ["process", "процесс"], ["thread", "поток"]],
    },
    {
        "slug": "networks_security",
        "title": "Сети и безопасность",
        "author": "",
        "topic": "computer_networks",
        "difficulty": "medium",
        "paragraphs": [
            "HTTP, DNS, TCP/IP и маршрутизация описывают сетевое взаимодействие.",
            "Информационная безопасность включает шифрование, аутентификацию и анализ уязвимостей.",
        ],
        "table": [["Термин", "Категория"], ["DNS", "network"], ["encryption", "security"]],
    },
    {
        "slug": "discrete_math",
        "title": "Дискретная математика",
        "author": "Synthetic NIR Dataset",
        "topic": "mathematics",
        "difficulty": "hard",
        "paragraphs": [
            "Граф, дерево, кратчайший путь, поток и паросочетание встречаются в алгоритмических задачах.",
            "Доказательство теоремы использует комбинаторику и оценку сложности.",
        ],
        "table": [["Объект", "Свойство"], ["graph", "вершины и ребра"], ["tree", "связный ациклический граф"]],
    },
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create synthetic DOCX samples for the NIR MVP dataset.")
    parser.add_argument("--output", default="data/real_input/docx", help="Output directory for DOCX files.")
    parser.add_argument("--manifest", default="datasets/manifest.csv", help="Dataset manifest path.")
    return parser.parse_args()


def read_manifest(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8", newline="") as file_obj:
        return list(csv.DictReader(file_obj))


def write_manifest(path: Path, rows: list[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as file_obj:
        writer = csv.DictWriter(file_obj, fieldnames=MANIFEST_FIELDS)
        writer.writeheader()
        writer.writerows(rows)


def upsert_manifest_rows(manifest_path: Path, new_rows: list[dict[str, str]]) -> None:
    rows = read_manifest(manifest_path)
    by_path = {row["relative_path"]: row for row in rows if row.get("relative_path")}
    for row in new_rows:
        by_path[row["relative_path"]] = row
    write_manifest(manifest_path, list(by_path.values()))


def create_docx(sample: dict[str, object], output_path: Path) -> None:
    document = Document()
    document.core_properties.title = str(sample["title"])
    if sample.get("author"):
        document.core_properties.author = str(sample["author"])
    document.core_properties.subject = "Synthetic educational material for NIR MVP"
    document.core_properties.keywords = str(sample["topic"])

    document.add_heading(str(sample["title"]), level=1)
    for paragraph in sample["paragraphs"]:  # type: ignore[index]
        document.add_paragraph(str(paragraph))

    table_data = sample["table"]  # type: ignore[index]
    table = document.add_table(rows=len(table_data), cols=2)
    for row_index, row_values in enumerate(table_data):
        for col_index, value in enumerate(row_values):
            table.cell(row_index, col_index).text = str(value)

    document.save(output_path)


def main() -> None:
    args = parse_args()
    output = Path(args.output)
    output.mkdir(parents=True, exist_ok=True)

    manifest_rows: list[dict[str, str]] = []
    for index, sample in enumerate(SAMPLES, start=1):
        file_name = f"synthetic_{index:02d}_{sample['slug']}.docx"
        output_path = output / file_name
        create_docx(sample, output_path)
        manifest_rows.append(
            {
                "file_id": f"docx_synth_{index:03d}",
                "file_name": file_name,
                "relative_path": str(output_path.as_posix()),
                "format": "docx",
                "source_type": "synthetic",
                "source_url": "",
                "topic_expected": str(sample["topic"]),
                "difficulty_expected": str(sample["difficulty"]),
                "has_text_layer": "true",
                "is_duplicate": "false",
                "is_bad_file": "false",
                "comment": "Synthetic DOCX sample with headings, paragraphs, table, and selected core properties.",
            }
        )
        print(f"[created] {output_path}")

    upsert_manifest_rows(Path(args.manifest), manifest_rows)
    print(f"Manifest updated: {args.manifest}")


if __name__ == "__main__":
    main()
