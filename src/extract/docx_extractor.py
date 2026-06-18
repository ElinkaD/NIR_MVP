"""DOCX extraction adapter based on python-docx."""

from __future__ import annotations

from pathlib import Path

from .base import ExtractionResult


class DocxExtractor:
    def extract(self, path: Path) -> ExtractionResult:
        try:
            from docx import Document
        except Exception as exc:  # pragma: no cover - depends on optional package state
            return ExtractionResult(
                extraction_status="failed",
                errors=[f"python-docx import failed: {exc}"],
            )

        try:
            document = Document(str(path))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            headings = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
                and paragraph.style is not None
                and paragraph.style.name.lower().startswith("heading")
            ]

            table_lines: list[str] = []
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_lines.append(" | ".join(cells))

            props = document.core_properties
            metadata = {
                "title": props.title or None,
                "author": props.author or None,
                "subject": props.subject or None,
                "keywords": props.keywords or None,
                "created": props.created.isoformat() if props.created else None,
                "modified": props.modified.isoformat() if props.modified else None,
            }
            text = "\n\n".join(paragraphs + table_lines)
            return ExtractionResult(
                text=text,
                page_count=None,
                title=metadata["title"],
                author=metadata["author"],
                headings=headings,
                metadata=metadata,
                extractor_used="python-docx",
                extractor_status="success" if text.strip() else "failed",
                extraction_status="success" if text.strip() else "failed",
                errors=[] if text.strip() else ["DOCX text is empty"],
            )
        except Exception as exc:
            return ExtractionResult(
                extractor_used="python-docx",
                extractor_status="failed",
                extraction_status="failed",
                errors=[f"DOCX extraction failed: {exc}"],
            )
